from __future__ import annotations

from io import BytesIO
from subprocess import CalledProcessError
from unittest.mock import Mock, patch

import pymupdf
from apps.policies.ai_service import (
    PolicyRecognitionError,
    _build_vision_messages,
    _prepare_image_for_vision,
    _prepare_pdf_page_for_vision,
    _render_pdf_pages_for_vision,
    _resolve_ai_client_config,
    extract_text_from_bytes,
    is_extracted_policy_text_poor,
    recognize_policy_from_bytes,
)
from django.test import SimpleTestCase, override_settings
from docx import Document
from PIL import Image


class ExtractPolicyTextFromBytesTests(SimpleTestCase):
    def test_docx_text_is_extracted(self):
        document = Document()
        document.add_paragraph("Полис КАСКО")
        document.add_paragraph("Номер ABC-123")
        buffer = BytesIO()
        document.save(buffer)

        text = extract_text_from_bytes(buffer.getvalue(), "policy.docx")

        self.assertIn("Полис КАСКО", text)
        self.assertIn("Номер ABC-123", text)

    @patch("apps.policies.ai_service.shutil.which", return_value=None)
    def test_doc_without_soffice_raises_clear_error(self, which_mock: Mock):
        with self.assertRaises(PolicyRecognitionError) as exc_info:
            extract_text_from_bytes(b"doc-bytes", "policy.doc")

        self.assertIn("LibreOffice", str(exc_info.exception))
        which_mock.assert_called_once_with("soffice")

    @patch("apps.policies.ai_service.Path.read_text", return_value="Полис DOC")
    @patch("apps.policies.ai_service.Path.exists", return_value=True)
    @patch("apps.policies.ai_service.subprocess.run")
    @patch("apps.policies.ai_service.shutil.which", return_value="/usr/bin/soffice")
    def test_doc_text_is_extracted_via_soffice(
        self,
        which_mock: Mock,
        run_mock: Mock,
        exists_mock: Mock,
        read_text_mock: Mock,
    ):
        text = extract_text_from_bytes(b"doc-bytes", "policy.doc")

        self.assertEqual(text, "Полис DOC")
        which_mock.assert_called_once_with("soffice")
        run_mock.assert_called_once()
        exists_mock.assert_called()
        read_text_mock.assert_called_once_with(encoding="utf-8", errors="ignore")

    @patch(
        "apps.policies.ai_service.subprocess.run",
        side_effect=CalledProcessError(1, ["soffice"]),
    )
    @patch("apps.policies.ai_service.shutil.which", return_value="/usr/bin/soffice")
    def test_doc_failed_conversion_raises_clear_error(
        self, which_mock: Mock, run_mock: Mock
    ):
        with self.assertRaises(PolicyRecognitionError) as exc_info:
            extract_text_from_bytes(b"doc-bytes", "policy.doc")

        self.assertIn("LibreOffice", str(exc_info.exception))
        which_mock.assert_called_once_with("soffice")
        run_mock.assert_called_once()


class PolicyVisionFallbackTests(SimpleTestCase):
    def _image_bytes(
        self, image_format: str = "PNG", *, orientation: int | None = None
    ):
        image = Image.new("RGB", (20, 10), "white")
        buffer = BytesIO()
        save_args = {}
        if orientation is not None:
            exif = Image.Exif()
            exif[274] = orientation
            save_args["exif"] = exif
        image.save(buffer, format=image_format, **save_args)
        return buffer.getvalue()

    def test_pdf_with_garbled_text_is_poor_text_candidate(self):
        text = "\x04\x05 \x06\x07 abc def ghijk " * 20

        self.assertTrue(is_extracted_policy_text_poor(text))

    def test_raw_pdf_stream_is_poor_text_candidate(self):
        text = "%PDF-1.7\n1 0 obj\n<< /Length 42 >>\nstream\npolicy\nendstream\n"

        self.assertTrue(is_extracted_policy_text_poor(text))

    @patch("apps.policies.ai_service.PdfReader", side_effect=Exception("bad pdf"))
    def test_pdf_parse_failure_does_not_fallback_to_raw_bytes(
        self, pdf_reader_mock: Mock
    ):
        with self.assertRaises(PolicyRecognitionError) as exc_info:
            extract_text_from_bytes(b"%PDF-1.7\nstream", "broken.pdf")

        self.assertIn("PDF", str(exc_info.exception))
        pdf_reader_mock.assert_called_once()

    @override_settings(POLICY_RECOGNITION_VISION_FALLBACK_ENABLED=True)
    @patch("apps.policies.ai_service.recognize_policy_from_pdf_images")
    @patch("apps.policies.ai_service.recognize_policy_from_text")
    @patch("apps.policies.ai_service.extract_text_from_bytes")
    def test_pdf_parse_failure_uses_vision_without_text_request(
        self,
        extract_mock: Mock,
        text_recognize_mock: Mock,
        vision_recognize_mock: Mock,
    ):
        extract_mock.side_effect = PolicyRecognitionError("PDF text extraction failed")
        expected = {"policy": {"policy_number": "SYS-1"}, "payments": []}
        vision_recognize_mock.return_value = (expected, "vision transcript")

        data, transcript = recognize_policy_from_bytes(b"%PDF", filename="policy.pdf")

        self.assertEqual(data, expected)
        self.assertEqual(transcript, "vision transcript")
        text_recognize_mock.assert_not_called()
        vision_recognize_mock.assert_called_once()

    @override_settings(POLICY_RECOGNITION_VISION_FALLBACK_ENABLED=True)
    @patch("apps.policies.ai_service.recognize_policy_from_pdf_images")
    @patch("apps.policies.ai_service.recognize_policy_from_text")
    @patch("apps.policies.ai_service.extract_text_from_bytes")
    def test_vision_failure_never_returns_weak_text_result(
        self,
        extract_mock: Mock,
        text_recognize_mock: Mock,
        vision_recognize_mock: Mock,
    ):
        extract_mock.return_value = (
            "Полис ОСАГО номер SYS-1 страховая премия автомобиль договор"
        )
        text_recognize_mock.return_value = ({"policy": {}, "payments": []}, "weak text")
        vision_recognize_mock.side_effect = PolicyRecognitionError("vision unavailable")

        with self.assertRaises(PolicyRecognitionError):
            recognize_policy_from_bytes(b"%PDF", filename="policy.pdf")

        text_recognize_mock.assert_called_once()
        vision_recognize_mock.assert_called_once()

    @override_settings(POLICY_RECOGNITION_VISION_FALLBACK_ENABLED=True)
    @patch("apps.policies.ai_service.recognize_policy_from_pdf_images")
    @patch("apps.policies.ai_service.recognize_policy_from_text")
    @patch("apps.policies.ai_service.extract_text_from_bytes")
    def test_invalid_text_mode_uses_vision_fallback(
        self,
        extract_mock: Mock,
        text_recognize_mock: Mock,
        vision_recognize_mock: Mock,
    ):
        extract_mock.return_value = (
            "Полис ОСАГО номер SYS-1 страховая премия автомобиль договор"
        )
        text_recognize_mock.side_effect = PolicyRecognitionError("bad json")
        expected = {"policy": {"policy_number": "SYS-1"}, "payments": []}
        vision_recognize_mock.return_value = (expected, "vision transcript")

        data, transcript = recognize_policy_from_bytes(
            b"%PDF",
            filename="policy.pdf",
        )

        self.assertEqual(data, expected)
        self.assertEqual(transcript, "vision transcript")
        vision_recognize_mock.assert_called_once()

    @override_settings(POLICY_RECOGNITION_VISION_FALLBACK_ENABLED=True)
    @patch("apps.policies.ai_service.recognize_policy_from_pdf_images")
    @patch("apps.policies.ai_service.recognize_policy_from_text")
    @patch("apps.policies.ai_service.extract_text_from_bytes")
    def test_empty_text_mode_result_uses_vision_fallback(
        self,
        extract_mock: Mock,
        text_recognize_mock: Mock,
        vision_recognize_mock: Mock,
    ):
        extract_mock.return_value = (
            "Полис ОСАГО номер SYS-1 страховая премия автомобиль договор"
        )
        text_recognize_mock.return_value = (
            {
                "client_name": "",
                "policy": {
                    "policy_number": "",
                    "insurance_type": "",
                    "insurance_company": "",
                    "contractor": "",
                    "sales_channel": "",
                    "start_date": "",
                    "end_date": "",
                    "vehicle_brand": "",
                    "vehicle_model": "",
                    "vehicle_vin": "",
                    "note": "импортировано с помощью ИИ",
                },
                "payments": [],
            },
            "text transcript",
        )
        expected = {"policy": {"policy_number": "SYS-1"}, "payments": []}
        vision_recognize_mock.return_value = (expected, "vision transcript")

        data, transcript = recognize_policy_from_bytes(
            b"%PDF",
            filename="policy.pdf",
        )

        self.assertEqual(data, expected)
        self.assertEqual(transcript, "vision transcript")
        vision_recognize_mock.assert_called_once()

    @override_settings(POLICY_RECOGNITION_VISION_FALLBACK_ENABLED=True)
    @patch("apps.policies.ai_service.recognize_policy_from_pdf_images")
    @patch("apps.policies.ai_service.recognize_policy_from_text")
    @patch("apps.policies.ai_service.extract_text_from_bytes")
    def test_good_text_mode_result_does_not_use_vision(
        self,
        extract_mock: Mock,
        text_recognize_mock: Mock,
        vision_recognize_mock: Mock,
    ):
        extract_mock.return_value = (
            "Полис ОСАГО номер SYS-1 страховая премия автомобиль договор"
        )
        expected = {
            "client_name": "Иванов Иван",
            "policy": {
                "policy_number": "SYS-1",
                "insurance_type": "ОСАГО",
                "insurance_company": "РЕСО",
                "contractor": "",
                "sales_channel": "",
                "start_date": "2026-01-01",
                "end_date": "2027-01-01",
                "vehicle_brand": "",
                "vehicle_model": "",
                "vehicle_vin": "",
                "note": "импортировано с помощью ИИ",
            },
            "payments": [
                {
                    "amount": 1000,
                    "payment_date": "2026-01-01",
                    "actual_payment_date": "2026-01-01",
                }
            ],
        }
        text_recognize_mock.return_value = (expected, "text transcript")

        data, transcript = recognize_policy_from_bytes(
            b"%PDF",
            filename="policy.pdf",
        )

        self.assertEqual(data, expected)
        self.assertEqual(transcript, "text transcript")
        vision_recognize_mock.assert_not_called()

    @override_settings(POLICY_RECOGNITION_MAX_VISION_PAGES=1)
    @patch("apps.policies.ai_service._detect_pdf_page_rotation", return_value=0)
    def test_pdf_render_limits_page_count(self, rotation_mock: Mock):
        document = pymupdf.open()
        document.new_page().insert_text((72, 72), "Страница 1")
        document.new_page().insert_text((72, 72), "Страница 2")
        content = document.tobytes()
        document.close()

        images = _render_pdf_pages_for_vision(content, "policy.pdf")

        self.assertEqual(len(images), 1)
        self.assertTrue(images[0].startswith(b"\xff\xd8"))
        rotation_mock.assert_called_once()

    @override_settings(
        POLICY_RECOGNITION_PDF_RENDER_DPI=72,
        POLICY_RECOGNITION_MAX_IMAGE_DIMENSION=100,
    )
    @patch("apps.policies.ai_service._detect_pdf_page_rotation", return_value=0)
    def test_rendered_pdf_page_is_resized_and_jpeg_normalized(
        self, rotation_mock: Mock
    ):
        document = pymupdf.open()
        document.new_page(width=400, height=200)
        content = document.tobytes()
        document.close()

        image_bytes = _render_pdf_pages_for_vision(content, "policy.pdf")[0]

        with Image.open(BytesIO(image_bytes)) as image:
            self.assertEqual(image.format, "JPEG")
            self.assertEqual(image.size, (100, 50))
        rotation_mock.assert_called_once()

    @patch(
        "apps.policies.ai_service.pytesseract.image_to_osd",
        return_value={"rotate": "90", "orientation_conf": "10.5"},
    )
    def test_pdf_page_osd_rotations_are_applied(self, osd_mock: Mock):
        expected_sizes = {90: (20, 40), 180: (40, 20), 270: (20, 40)}
        for rotation, expected_size in expected_sizes.items():
            with self.subTest(rotation=rotation):
                source = BytesIO()
                Image.new("RGB", (40, 20), "white").save(source, format="PNG")
                osd_mock.return_value = {
                    "rotate": str(rotation),
                    "orientation_conf": "10.5",
                }

                normalized = _prepare_pdf_page_for_vision(source.getvalue())

                with Image.open(BytesIO(normalized)) as image:
                    self.assertEqual(image.size, expected_size)
        self.assertEqual(osd_mock.call_count, 3)

    @patch(
        "apps.policies.ai_service.pytesseract.image_to_osd",
        return_value={"rotate": "90", "orientation_conf": "4.9"},
    )
    def test_pdf_page_osd_low_confidence_keeps_orientation(self, osd_mock: Mock):
        source = BytesIO()
        Image.new("RGB", (40, 20), "white").save(source, format="PNG")

        normalized = _prepare_pdf_page_for_vision(source.getvalue())

        with Image.open(BytesIO(normalized)) as image:
            self.assertEqual(image.size, (40, 20))
        osd_mock.assert_called_once()

    @patch(
        "apps.policies.ai_service.pytesseract.image_to_osd",
        side_effect=RuntimeError("OSD unavailable"),
    )
    def test_pdf_page_osd_failure_keeps_orientation(self, osd_mock: Mock):
        source = BytesIO()
        Image.new("RGB", (40, 20), "white").save(source, format="PNG")

        normalized = _prepare_pdf_page_for_vision(source.getvalue())

        with Image.open(BytesIO(normalized)) as image:
            self.assertEqual(image.size, (40, 20))
        osd_mock.assert_called_once()

    @override_settings(POLICY_RECOGNITION_MAX_IMAGE_DIMENSION=2048)
    def test_jpeg_is_normalized_and_exif_orientation_is_applied(self):
        normalized = _prepare_image_for_vision(
            self._image_bytes("JPEG", orientation=6), "policy.jpg"
        )

        with Image.open(BytesIO(normalized)) as image:
            self.assertEqual(image.format, "JPEG")
            self.assertEqual(image.size, (10, 20))

    def test_png_is_normalized_for_vision(self):
        normalized = _prepare_image_for_vision(self._image_bytes(), "policy.png")

        with Image.open(BytesIO(normalized)) as image:
            self.assertEqual(image.format, "JPEG")
            self.assertEqual(image.size, (20, 10))

    @override_settings(POLICY_RECOGNITION_MAX_IMAGE_DIMENSION=100)
    def test_image_is_resized_to_configured_max_dimension(self):
        image = Image.new("RGB", (400, 200), "white")
        source = BytesIO()
        image.save(source, format="JPEG")

        normalized = _prepare_image_for_vision(source.getvalue(), "policy.jpg")

        with Image.open(BytesIO(normalized)) as prepared:
            self.assertEqual(prepared.size, (100, 50))

    def test_image_with_alpha_stays_png(self):
        image = Image.new("RGBA", (20, 10), (255, 255, 255, 0))
        source = BytesIO()
        image.save(source, format="PNG")

        normalized = _prepare_image_for_vision(source.getvalue(), "policy.png")

        self.assertTrue(normalized.startswith(b"\x89PNG"))

    @override_settings(
        POLICY_RECOGNITION_MODEL="google/gemini-2.5-flash-lite",
        AI_MODEL="gpt-4o-mini",
        AI_API_KEY="test-key",  # pragma: allowlist secret
    )
    def test_policy_model_has_priority_over_common_model(self):
        _, _, model = _resolve_ai_client_config()

        self.assertEqual(model, "google/gemini-2.5-flash-lite")

    @override_settings(
        POLICY_RECOGNITION_MODEL="",
        AI_MODEL="gpt-4o-mini",
        AI_API_KEY="test-key",  # pragma: allowlist secret
    )
    def test_policy_model_falls_back_to_common_model(self):
        _, _, model = _resolve_ai_client_config()

        self.assertEqual(model, "gpt-4o-mini")

    @override_settings(
        AI_API_KEY="test-key",  # pragma: allowlist secret
        AI_BASE_URL="",
        AI_MODEL="",
        POLICY_RECOGNITION_MODEL="",
    )
    def test_ai_client_config_uses_polza_defaults(self):
        _, base_url, model = _resolve_ai_client_config()

        self.assertEqual(base_url, "https://polza.ai/api/v1")
        self.assertEqual(model, "google/gemini-2.5-flash-lite")

    @override_settings(AI_API_KEY="")
    def test_ai_client_config_requires_api_key(self):
        with self.assertRaisesMessage(ValueError, "AI_API_KEY не задан"):
            _resolve_ai_client_config()

    def test_broken_image_raises_clear_error(self):
        with self.assertRaises(PolicyRecognitionError) as exc_info:
            _prepare_image_for_vision(b"not-an-image", "broken.png")

        self.assertIn("Не удалось подготовить изображение", str(exc_info.exception))

    def test_unsupported_image_format_raises_clear_error(self):
        with self.assertRaises(PolicyRecognitionError) as exc_info:
            _prepare_image_for_vision(b"image", "policy.webp")

        self.assertIn("Неподдерживаемый формат", str(exc_info.exception))

    @override_settings(POLICY_RECOGNITION_MAX_VISION_PAGES=1)
    def test_pdf_and_images_share_visual_input_limit(self):
        document = pymupdf.open()
        document.new_page().insert_text((72, 72), "Page 1")
        pdf_content = document.tobytes()
        document.close()

        messages, _ = _build_vision_messages(
            [
                {"name": "first.png", "content": self._image_bytes()},
                {"name": "second.pdf", "content": pdf_content},
            ]
        )

        user_content = messages[1]["content"]
        self.assertEqual(sum(item["type"] == "image_url" for item in user_content), 1)
